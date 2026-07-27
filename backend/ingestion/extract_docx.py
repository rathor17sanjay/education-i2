"""Word document (.docx) text extraction. Same extract(path) -> str
contract as extract_pdf.py, so the admin upload dispatcher can treat every
source kind uniformly.
"""

from docx import Document


def extract(path: str) -> str:
    doc = Document(path)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Admissions docs often carry fee/eligibility tables -- extracting these
    # separately means that data isn't silently dropped, since python-docx
    # doesn't interleave table text into the paragraph stream automatically.
    table_texts = []
    for table in doc.tables:
        rows = ["\t".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        table_texts.append("\n".join(rows))

    sections = ["\n\n".join(paragraphs)]
    sections.extend(table_texts)
    return "\n\n".join(s for s in sections if s.strip())
